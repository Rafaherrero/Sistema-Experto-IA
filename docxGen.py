#!/usr/bin/env python

#install
#pip install python-docx or
#download from http://pypi.python.org/pypi/python-docx
#and python setup.py install
from docx import Document
from docx.shared import Inches
document = Document()
animales = ["Abeja","Acaro","Aguila","Arania","Ardilla","Armadillo","Asno","Atun","Ballena","Beluga",
			"Bufalo","Buho","Caballito de mar","Caballo","Cabra","Calamar","Camaleon" "Camello","Cangrejo",
    		"Canguro","Caracol","Castor","Cebra","Cerdo","Chimpance","Ciempies","Ciervo","Ciguenia","Cisne",
    		"Cocodrilo","Codorniz","Colibri","Conejo","Coyote","Cucaracha","Cuervo","Delfin","Dromedario",
    		"Elefante","Erizo","Escarabajo","Escorpion","Foca","Gallina","Gato","Gorila","Grillo","Hamster",
    		"Hipopotamo","Hormiga","Jabali","Jirafa","Koala","Lagarto","Lemur","Leon","Lince" "Llama","Lobo",
    		"Lombriz","Loro","Mandril","Mariposa","Mariquita","Medusa","Mosca","Mosquito","Murcielago","Narval",
    		"Oca","Orangutan","Ornitorrinco","Oruga","Oso","Oveja","Paloma","Pantera","Pato","Pavo","Perro",
    		"Pez Espada","Pinguinos","Pirania","Polilla","Pulpo","Rana","Rata","Rinoceronte","Salamandra",
    		"Saltamontes","Sepia","Serpiente","Suricato","Tiburon","Tigre","Topo","Toro","Tortuga","Vaca","Zorro"
]

document.add_heading('Representacion del Conocimiento en OVA', 0)
sistema=open("Archivos_finales/animales.clp","r")
while True:
	li = sistema.readline()
	if not li: break
	for x in range(len(animales)):
		if "(defrule "+animales[x] in li:
			table = document.add_table(rows=1, cols=3)
			table.style = 'TableGrid'
			table.rows[0].style = "background-color:gray"
			hdr_cells = table.rows[0].cells
			hdr_cells[0].text = 'Objeto'
			hdr_cells[1].text = 'Atributo'
			hdr_cells[2].text = 'Valor'
			while int(len(li))!=1:
				li=sistema.readline()
				OVA=li.split()
				#print len(OVA)
				if len(OVA)==2:
					#print li
					row_cells = table.add_row().cells
					row_cells[0].text = animales[x]
					row_cells[1].text = OVA[0][1:int(len(OVA[0]))]
					row_cells[2].text = OVA[1][0:int(len(OVA[1]))-1]
					print "Cargado "+animales[x]+" "+OVA[0][1:int(len(OVA[0]))]+" "+OVA[1][0:int(len(OVA[1]))-1]
					
				if len(OVA)==4:
					#print li
					if OVA[3][0:int(len(OVA[3]))-1]!=animales[x]:
						row_cells = table.add_row().cells
						row_cells[0].text = animales[x]
						row_cells[1].text = OVA[2]
						row_cells[2].text = OVA[3][0:int(len(OVA[3]))-1]
						print "Cargado "+animales[x]+" "+OVA[2]+" "+OVA[3][0:int(len(OVA[3]))-1]
						document.add_paragraph()

			#if li==')':break
				#document.add_page_break()





document.save('Esquema.docx')
print "Creado el Esquema del Sistema  (Esquema.docx)"